"""Parse DOCX files into structural elements."""

from __future__ import annotations

from pathlib import Path

from docx import Document

from .elements import Element, HeadingElement, TableRowElement, TextElement


def _is_heading_style(style_name: str) -> bool:
    if not style_name:
        return False
    name = style_name.lower()
    return name.startswith("heading") or name == "title" or name == "subtitle"


def _heading_level(style_name: str) -> int:
    name = style_name.lower()
    if name == "title":
        return 1
    if name == "subtitle":
        return 2
    for ch in name:
        if ch.isdigit():
            return int(ch)
    return 1


def _is_section_break_row(cells: list[str]) -> bool:
    """A row where only the first cell has content is likely a section header."""
    non_empty = [c for c in cells if c.strip()]
    return len(non_empty) == 1 and cells[0].strip() != ""


def parse_docx(path: Path) -> list[Element]:
    doc = Document(path)
    elements: list[Element] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        style_name = para.style.name if para.style else ""

        if _is_heading_style(style_name):
            elements.append(
                HeadingElement(text=text, level=_heading_level(style_name))
            )
        else:
            elements.append(TextElement(text=text))

    for table_idx, table in enumerate(doc.tables):
        headers: list[str] = []

        for row_idx, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]

            if row_idx == 0:
                headers = cells
                continue

            if not any(cells):
                continue

            is_section_break = _is_section_break_row(cells)

            elements.append(
                TableRowElement(
                    cells=cells,
                    headers=headers,
                    is_section_break=is_section_break,
                    table_index=table_idx,
                    row_index=row_idx,
                )
            )

    return elements
